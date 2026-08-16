import json
import os
import logging
import traceback
from firebase_functions import https_fn, options, firestore_fn

logger = logging.getLogger(__name__)

from core.config import global_cors, GEMINI_MODEL

@firestore_fn.on_document_written(document="users/{uid}")
def sync_user_claims_on_write(event: firestore_fn.Event[firestore_fn.Change[firestore_fn.DocumentSnapshot]]) -> None:
    """
    Firestore trigger to synchronize user roles and cartorio_id to Auth Custom Claims.
    """
    from core.firebase_utils import _init_firebase
    _init_firebase()
    from firebase_admin import auth
    import logging

    logger = logging.getLogger(__name__)

    user_uid = event.params["uid"]

    # If the document is deleted, we don't necessarily delete the auth user here,
    # but we could clear claims.
    if not event.data.after:
        try:
            auth.set_custom_user_claims(user_uid, None)
            logger.info(f"Cleared custom claims for deleted user document: {user_uid}")
        except Exception as e:
            logger.error(f"Failed to clear custom claims for deleted user {user_uid}: {e}")
        return

    after_data = event.data.after.to_dict()

    # Extract fields we care about
    status = after_data.get("status")
    role = after_data.get("role")
    cartorio_id = after_data.get("cartorio_id")

    custom_claims = {}

    if status == "revoked":
        # Actively clear claims if revoked
        custom_claims = {"role": None, "cartorio_id": None}
    else:
        custom_claims = {"role": role, "cartorio_id": cartorio_id}

    try:
        auth.set_custom_user_claims(user_uid, custom_claims)
        logger.info(f"Successfully synced custom claims for user {user_uid}: {custom_claims}")
    except Exception as e:
        logger.error(f"Failed to sync custom claims for user {user_uid}: {e}")

@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.GB_2, timeout_sec=1200)
def extract_batch_document_data(req: https_fn.Request) -> https_fn.Response:
    """
    Extracts structured data autonomously from a batch of documents stored in GCS.
    Accepts POST requests with JSON payload: {"gcs_uris": ["gs://...", "gs://..."]}
    """
    if req.method != "POST":
        return https_fn.Response(
            json.dumps({"error": "Only POST requests are accepted"}),
            status=405,
            content_type="application/json"
        )

    try:
        data = req.get_json(silent=True)
        if not data:
            return https_fn.Response(
                json.dumps({"error": "Missing or malformed JSON payload"}),
                status=400,
                content_type="application/json"
            )

        gcs_uris = data.get("gcs_uris")

        if not gcs_uris or not isinstance(gcs_uris, list):
            return https_fn.Response(
                json.dumps({"error": "Missing or invalid gcs_uris (must be a list)"}),
                status=400,
                content_type="application/json"
            )

        from core.extractor import DocumentExtractor
        extractor = DocumentExtractor()
        extracted_data = extractor.extract_batch(gcs_uris)

        return https_fn.Response(
            json.dumps({"status": "success", "data": extracted_data}, ensure_ascii=False, allow_nan=False),
            status=200,
            content_type="application/json"
        )
    except ValueError as e:
        logger.error("ValueError in extract_batch_document_data", exc_info=True)
        return https_fn.Response(
            json.dumps({
                "error": str(e),
                "traceback": traceback.format_exc()
            }),
            status=400,
            content_type="application/json"
        )
    except Exception as e:
        logger.error("Error in extract_batch_document_data", exc_info=True)
        status_code = 500
        error_code = "INTERNAL_ERROR"
        if "429" in str(e) or "ResourceExhausted" in str(e) or "quota" in str(e).lower():
            status_code = 429
            error_code = "RESOURCE_EXHAUSTED"
        return https_fn.Response(
            json.dumps({
                "error": f"Internal server error: {str(e)}",
                "code": error_code,
                "traceback": traceback.format_exc()
            }),
            status=status_code,
            content_type="application/json"
        )


@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.GB_2, timeout_sec=1200)
def extract_document_data(req: https_fn.Request) -> https_fn.Response:
    """
    Extracts structured data from a document stored in GCS.
    Accepts POST requests with JSON payload: {"gcs_uri": "gs://...", "document_type": "..."}
    """
    if req.method != "POST":
        return https_fn.Response(
            json.dumps({"error": "Only POST requests are accepted"}),
            status=405,
            content_type="application/json"
        )

    try:
        data = req.get_json()
        if not data:
            return https_fn.Response(
                json.dumps({"error": "Missing JSON payload"}),
                status=400,
                content_type="application/json"
            )

        gcs_uri = data.get("gcs_uri")

        if not gcs_uri or not isinstance(gcs_uri, str):
            return https_fn.Response(
                json.dumps({"error": "Missing or invalid gcs_uri"}),
                status=400,
                content_type="application/json"
            )

        document_type = data.get("document_type")
        minuta_id = data.get("minuta_id")

        from core.extractor import DocumentExtractor
        extractor = DocumentExtractor()
        extracted_data = extractor.extract(gcs_uri, document_type=document_type)

        if minuta_id:
            from core.firebase_utils import _init_firebase
            _init_firebase()
            from firebase_admin import firestore
            db = firestore.client()
            minuta_ref = db.collection("minutas").document(minuta_id)
            try:
                minuta_ref.update({
                    "status": "hitl_required",
                    "ai_extracted_data": extracted_data,
                    "updatedAt": firestore.SERVER_TIMESTAMP
                })
            except Exception as update_err:
                logger.error(f"Failed to update minuta {minuta_id}: {update_err}")

        return https_fn.Response(
            json.dumps({"status": "success", "data": extracted_data}, ensure_ascii=False, allow_nan=False),
            status=200,
            content_type="application/json"
        )
    except ValueError as e:
        # Catch specific value errors raised during extraction (e.g., config missing)
        logger.error("ValueError in extract_document_data", exc_info=True)
        if req.get_json(silent=True) and req.get_json(silent=True).get("minuta_id"):
            minuta_id = req.get_json(silent=True).get("minuta_id")
            from core.firebase_utils import _init_firebase
            _init_firebase()
            from firebase_admin import firestore
            db = firestore.client()
            db.collection("minutas").document(minuta_id).update({"status": "error", "error": str(e)})

        return https_fn.Response(
            json.dumps({
                "error": str(e),
                "traceback": traceback.format_exc()
            }),
            status=400,
            content_type="application/json"
        )
    except Exception as e:
        logger.error("Error in extract_document_data", exc_info=True)
        if req.get_json(silent=True) and req.get_json(silent=True).get("minuta_id"):
            minuta_id = req.get_json(silent=True).get("minuta_id")
            from core.firebase_utils import _init_firebase
            _init_firebase()
            from firebase_admin import firestore
            db = firestore.client()
            db.collection("minutas").document(minuta_id).update({"status": "error", "error": f"Internal server error: {str(e)}"})

        status_code = 500
        error_code = "INTERNAL_ERROR"
        if "429" in str(e) or "ResourceExhausted" in str(e) or "quota" in str(e).lower():
            status_code = 429
            error_code = "RESOURCE_EXHAUSTED"
        return https_fn.Response(
            json.dumps({
                "error": f"Internal server error: {str(e)}",
                "code": error_code,
                "traceback": traceback.format_exc()
            }),
            status=status_code,
            content_type="application/json"
        )

@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.MB_256)
def submit_audit_event(req: https_fn.Request) -> https_fn.Response:
    """
    Accepts feedback on document validation and logs it asynchronously to Firestore.
    Accepts POST requests with JSON payload containing:
    - document_id (optional, string)
    - document_type (string)
    - ai_detected (dict)
    - user_corrected (dict)
    - validation_errors (list)
    """
    if req.method != "POST":
        return https_fn.Response(
            json.dumps({"error": "Only POST requests are accepted"}),
            status=405,
            content_type="application/json"
        )

    try:
        data = req.get_json()
        if not data:
            return https_fn.Response(
                json.dumps({"error": "Missing JSON payload"}),
                status=400,
                content_type="application/json"
            )

        # Validate required fields
        if "document_type" not in data or "ai_detected" not in data or "user_corrected" not in data:
            return https_fn.Response(
                json.dumps({"error": "Missing required fields: document_type, ai_detected, user_corrected"}),
                status=400,
                content_type="application/json"
            )

        if not isinstance(data.get("ai_detected"), dict) or not isinstance(data.get("user_corrected"), dict):
            return https_fn.Response(
                json.dumps({"error": "ai_detected and user_corrected must be dictionaries"}),
                status=400,
                content_type="application/json"
            )

        event_data = {
            "document_id": data.get("document_id", "unknown"),
            "document_type": data.get("document_type"),
            "ai_detected": data.get("ai_detected"),
            "user_corrected": data.get("user_corrected"),
            "validation_errors": data.get("validation_errors", [])
        }

        # Asynchronously log the event to Firestore
        from core.audit import log_audit_event_async
        log_audit_event_async(event_data)

        # Return 200 OK immediately
        return https_fn.Response(
            json.dumps({"status": "success", "message": "Audit event submitted"}),
            status=200,
            content_type="application/json"
        )
    except Exception as e:
        return https_fn.Response(
            json.dumps({"error": f"Internal server error: {str(e)}"}),
            status=500,
            content_type="application/json"
        )

from admin import inviteEmployee, revokeEmployeeAccess, reactivateEmployeeAccess

@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.MB_256)
def finalize_validation(req: https_fn.Request) -> https_fn.Response:
    """Finalizes the validation process and marks the draft as completed."""
    if req.method != "POST":
        return https_fn.Response(json.dumps({"error": "Method not allowed"}), status=405)

    try:
        data = req.get_json()
        if not data or not data.get("document_id") or not data.get("final_text"):
            return https_fn.Response(json.dumps({"error": "Missing document_id or final_text"}), status=400)

        document_id = data.get("document_id")
        final_text = data.get("final_text")

        auth_header = req.headers.get("Authorization")
        if not auth_header or not auth_header.startswith("Bearer "):
            return https_fn.Response(json.dumps({"error": "Unauthorized"}), status=401)

        token = auth_header.split("Bearer ")[1]

        cartorio_id = req.headers.get("X-Cartorio-ID")
        if not cartorio_id:
             return https_fn.Response(json.dumps({"error": "Missing X-Cartorio-ID header"}), status=400)

        _init_firebase()
        decoded_token = auth.verify_id_token(token)
        uid = decoded_token.get("uid")

        token_cartorio_id = decoded_token.get("cartorio_id")
        token_role = decoded_token.get("role")

        if token_role != "super_admin" and token_cartorio_id != cartorio_id:
             return https_fn.Response(json.dumps({"error": "Forbidden: Cannot access another tenant"}), status=403)

        db = firestore.client()
        minuta_ref = db.collection("minutas").document(document_id)
        minuta_doc = minuta_ref.get()

        if not minuta_doc.exists:
             return https_fn.Response(json.dumps({"error": "Document not found"}), status=404)

        minuta_data = minuta_doc.to_dict()
        if minuta_data.get("cartorio_id") != cartorio_id:
            return https_fn.Response(json.dumps({"error": "Forbidden"}), status=403)

        minuta_ref.update({
             "status": "completed",
             "human_final_data": {"final_text": final_text},
             "updatedAt": firestore.SERVER_TIMESTAMP,
             "finalizedBy": uid
        })

        return https_fn.Response(json.dumps({"status": "success"}), status=200, content_type="application/json")
    except Exception as e:
        logger.error(f"Error in finalize_validation: {e}", exc_info=True)
        return https_fn.Response(json.dumps({"error": str(e)}), status=500)


@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.MB_512, timeout_sec=120)
def orchestrate_document(req: https_fn.Request) -> https_fn.Response:
    """
    Phase 2 Orchestrator Endpoint.
    Uses semantic embedding, vector search, and LLM to assemble document clauses based on user intent.
    """
    if req.method != "POST":
        return https_fn.Response(json.dumps({"error": "Method not allowed"}), status=405, content_type="application/json")

    auth_header = req.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        return https_fn.Response(json.dumps({"error": "Unauthenticated"}), status=401, content_type="application/json")

    token = auth_header.split("Bearer ")[1]

    try:
        from core.firebase_utils import _init_firebase
        _init_firebase()
        from firebase_admin import auth, firestore
        from google.cloud.firestore_v1.vector import Vector
        from google.cloud.firestore_v1.base_vector_query import DistanceMeasure

        try:
            decoded_token = auth.verify_id_token(token)
            user_cartorio = decoded_token.get("cartorio_id")
            user_role = decoded_token.get("role")
        except Exception:
            return https_fn.Response(json.dumps({"error": "Invalid token"}), status=401, content_type="application/json")

        req_data = req.get_json(silent=True) or {}
        intent = req_data.get("intent")
        entities = req_data.get("entities", [])

        if not intent:
            return https_fn.Response(json.dumps({"error": "Missing intent"}), status=400, content_type="application/json")

        from core.generator import vectorize_text
        intent_vector = vectorize_text(intent)

        if not intent_vector:
            return https_fn.Response(json.dumps({"error": "Failed to vectorize intent"}), status=500, content_type="application/json")

        db = firestore.client()
        clauses_ref = db.collection("clauses")

        # KNN vector search on Firestore
        vector_query = clauses_ref.find_nearest(
            vector_field="embedding",
            query_vector=Vector(intent_vector),
            distance_measure=DistanceMeasure.COSINE,
            limit=10,
        )

        results = vector_query.get()

        candidates = []
        candidate_ids = set()
        for doc in results:
            doc_data = doc.to_dict()
            c_cartorio = doc_data.get("cartorio_id", "SYSTEM")
            # Enforce tenant isolation logic
            if user_role != "super_admin" and c_cartorio not in [user_cartorio, "SYSTEM"]:
                continue

            candidate_ids.add(doc.id)
            candidates.append({
                "id": doc.id,
                "title": doc_data.get("title", ""),
                "text": doc_data.get("text", ""),
                "required_variables": doc_data.get("required_variables", [])
            })

        if not candidates:
             return https_fn.Response(json.dumps({
                "selected_clause_ids": [],
                "reasoning": "No relevant clauses found."
             }), status=200, content_type="application/json")

        from google import genai
        from google.genai import types
        import os
        project_id = os.environ.get("FIREBASE_PROJECT_ID", "cartorio-meneghel-ai")
        location = os.environ.get("VERTEX_AI_LOCATION", "us-central1")
        client = genai.Client(vertexai=True, project=project_id, location=location)

        prompt = f"""
You are an expert legal orchestrator for a Brazilian Cartório.
The user wants to generate a document based on the following intent:
INTENT: "{intent}"

Below is a list of candidate clauses retrieved from our legal database:
{json.dumps(candidates, ensure_ascii=False, indent=2)}

Below are the extracted entities available from the user's uploaded document:
{json.dumps(entities, ensure_ascii=False, indent=2)}

Your task is twofold:
1. Select the optimal clauses from the candidates to fulfill the user's intent.
   Sequence them in the correct logical order for a legal document.
   ALWAYS start the sequence with a Preamble/Qualification clause (e.g., "Qualificação") that qualifies the involved parties (e.g., Outorgante and Outorgada) before listing the specific clauses.
   Do NOT invent new clauses. Only select from the provided candidates.
2. Based on the requested intent and selected clauses, map the extracted entities to logical abstract roles (e.g., OUTORGANTE, PROCURADOR, OUTORGADO, COMPRADOR, VENDEDOR).
   Return this mapping in the `role_mapping` field.
   The keys should be the logical abstract roles.
   The values should be an ARRAY of the entity IDs that belong to that role.
   IMPORTANT: You MUST use ONLY the exact `id` string (e.g., 'ent_123') in the `role_mapping` JSON. NEVER use the entity's name.
"""

        from pydantic import BaseModel, Field
        from typing import List, Dict

        class OrchestrationResponse(BaseModel):
            selected_clause_ids: List[str] = Field(description="List of selected clause IDs in logical sequence")
            role_mapping: Dict[str, List[str]] = Field(default_factory=dict, description="Dictionary mapping abstract roles (e.g., 'OUTORGANTE') to arrays of entity IDs (e.g., ['ent_123', 'ent_456']).")
            reasoning: str = Field(description="Explanation of why these clauses and roles were selected")

        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=OrchestrationResponse,
                temperature=0.0
            ),
        )

        parsed_response = json.loads(response.text)
        llm_selected_ids = parsed_response.get("selected_clause_ids", [])

        # Robust intersection check to prevent hallucination
        final_selected_ids = [cid for cid in llm_selected_ids if cid in candidate_ids]
        parsed_response["selected_clause_ids"] = final_selected_ids

        # Aggregate required variables from the selected candidates
        required_variables_map = {}
        for candidate in candidates:
            if candidate["id"] in final_selected_ids:
                for var in candidate.get("required_variables", []):
                    # Dedup by variable name
                    var_name = var.get("name")
                    if var_name and var_name not in required_variables_map:
                        required_variables_map[var_name] = var

        parsed_response["required_variables"] = list(required_variables_map.values())

        return https_fn.Response(
            json.dumps(parsed_response),
            status=200,
            content_type="application/json"
        )

    except Exception as e:
        logger.error(f"Error in orchestrate_document: {e}", exc_info=True)
        return https_fn.Response(json.dumps({"error": str(e)}), status=500, content_type="application/json")

@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.MB_256)
def api_status(req: https_fn.Request) -> https_fn.Response:
    """Returns the API status."""
    return https_fn.Response(
        json.dumps({"status": "online", "version": "1.0.0"}),
        content_type="application/json"
    )

@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.MB_256)
def merge_entities(req: https_fn.Request) -> https_fn.Response:
    """
    Merges an array of entities using the backend Master Truth Profile engine.
    Accepts POST requests with JSON payload: {"entities": [...]}
    """
    if req.method != "POST":
        return https_fn.Response(
            json.dumps({"error": "Only POST requests are accepted"}),
            status=405,
            content_type="application/json"
        )

    try:
        data = req.get_json()
        if not data:
            return https_fn.Response(
                json.dumps({"error": "Missing JSON payload"}),
                status=400,
                content_type="application/json"
            )

        entities = data.get("entities")

        if not entities or not isinstance(entities, list):
            return https_fn.Response(
                json.dumps({"error": "Missing or invalid entities (must be a list)"}),
                status=400,
                content_type="application/json"
            )

        from core.consolidator import MasterProfileConsolidator
        from core.models import validate_entity
        merged_entities = MasterProfileConsolidator.deduplicate_entities(entities)
        validated_entities = [validate_entity(ent) for ent in merged_entities]

        return https_fn.Response(
            json.dumps({"status": "success", "entities": validated_entities}),
            status=200,
            content_type="application/json"
        )
    except Exception as e:
        logger.error("Error in merge_entities", exc_info=True)
        return https_fn.Response(
            json.dumps({"error": f"Internal server error: {str(e)}"}),
            status=500,
            content_type="application/json"
        )

@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.MB_512, timeout_sec=540)
def format_draft(req: https_fn.Request) -> https_fn.Response:
    """
    Formats the raw draft text by securely injecting ground truth entities using an LLM.
    Accepts POST requests with JSON payload: {"raw_text": "...", "ground_truth": {...}}
    """
    if req.method != "POST":
        return https_fn.Response(
            json.dumps({"error": "Only POST requests are accepted"}),
            status=405,
            content_type="application/json"
        )

    try:
        data = req.get_json()
        if not data:
            return https_fn.Response(
                json.dumps({"error": "Missing JSON payload"}),
                status=400,
                content_type="application/json"
            )

        raw_text = data.get("raw_text")
        ground_truth = data.get("ground_truth")

        if not raw_text or ground_truth is None:
             return https_fn.Response(
                json.dumps({"error": "Missing required fields: raw_text or ground_truth"}),
                status=400,
                content_type="application/json"
            )

        if not isinstance(raw_text, str) or not isinstance(ground_truth, dict):
            return https_fn.Response(
                json.dumps({"error": "Invalid payload types. Expected string for raw_text and dict for ground_truth"}),
                status=400,
                content_type="application/json"
            )

        import re
        footer_text = ""
        # Match "Emolumentos:" explicitly to prevent prematurely cutting off the document
        # on standard legal clauses that happen to start with the word "Emolumentos"
        footer_match = re.search(r'(?im)^(\s*emolumentos:\s*R\$[\s\S]*)', raw_text)
        if footer_match:
            footer_text = footer_match.group(1)
            raw_text = raw_text[:footer_match.start()]

        from google import genai
        from google.genai import types

        # Initialize the Gemini client
        project_id = os.environ.get("FIREBASE_PROJECT_ID", "cartorio-meneghel-ai")
        location = os.environ.get("VERTEX_AI_LOCATION", "us-central1")

        client = genai.Client(vertexai=True, project=project_id, location=location)

        prompt = f"""
<TASK>
Update the following raw draft text by securely injecting the provided ground truth entity data.
</TASK>

<BUSINESS_RULES>
1. You must act STRICTLY as a grammar-aware data-entry formatter.
2. Inject the facts from the ground truth JSON into the text.
3. Conjugate Portuguese grammar correctly (e.g., feminine vs masculine, singular vs plural).
4. For missing fields in the text, insert the data naturally.
5. DO NOT alter any legal boilerplate, punctuation, or formatting outside of the entity data.
CRITICAL: If the Ground Truth JSON contains fields (like RG or naturalidade) that are missing in the raw text, you MUST insert them naturally into the sentence (e.g., 'portador do RG X', 'natural de Y'). Do not drop data.
</BUSINESS_RULES>

<GROUND_TRUTH_JSON>
{json.dumps(ground_truth, ensure_ascii=False, indent=2)}
</GROUND_TRUTH_JSON>

<RAW_DRAFT_TEXT>
{raw_text}
</RAW_DRAFT_TEXT>

Return the updated final text ONLY, without any markdown formatting or explanations.
"""
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.0
            )
        )

        formatted_text = response.text.strip()
        # Remove potential markdown code blocks if the LLM adds them despite instructions
        if formatted_text.startswith("```"):
            lines = formatted_text.split('\n')
            if len(lines) > 2:
                formatted_text = '\n'.join(lines[1:-1])

        if footer_text:
            formatted_text = formatted_text.strip() + "\n\n" + footer_text.strip()

        return https_fn.Response(
            json.dumps({"status": "success", "formatted_text": formatted_text}),
            status=200,
            content_type="application/json"
        )
    except Exception as e:
        logger.error("Error in format_draft", exc_info=True)
        return https_fn.Response(
            json.dumps({"error": f"Internal server error: {str(e)}"}),
            status=500,
            content_type="application/json"
        )


@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.MB_512, timeout_sec=540)
def validate_document_text(req: https_fn.Request) -> https_fn.Response:
    """
    Validates typed text against ground truth deterministically.
    Accepts POST requests with JSON payload: {"ground_truth": {...}, "typed_text": "..."}
    """
    if req.method != "POST":
        return https_fn.Response(
            json.dumps({"error": "Only POST requests are accepted"}),
            status=405,
            content_type="application/json"
        )

    try:
        data = req.get_json()
        if not data:
            return https_fn.Response(
                json.dumps({"error": "Missing JSON payload"}),
                status=400,
                content_type="application/json"
            )

        ground_truth = data.get("ground_truth", {})
        typed_text = data.get("typed_text", "")

        if not isinstance(ground_truth, dict) or not isinstance(typed_text, str):
            return https_fn.Response(
                json.dumps({"error": "Invalid payload types. Expected dict for ground_truth and string for typed_text"}),
                status=400,
                content_type="application/json"
            )

        from core.validator import DocumentValidator
        validator = DocumentValidator(ground_truth, typed_text)
        errors = validator.validate()

        return https_fn.Response(
            json.dumps({"status": "success", "errors": errors}),
            status=200,
            content_type="application/json"
        )
    except Exception as e:
        return https_fn.Response(
            json.dumps({"error": f"Internal server error: {str(e)}"}),
            status=500,
            content_type="application/json"
        )
@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.MB_256)
def log_audit_event(req: https_fn.Request) -> https_fn.Response:
    """
    Logs an audit event, such as marking a document as unreadable.
    Accepts POST requests with JSON payload: {"file_name": "...", "quality_flag": true/false}
    """
    if req.method != "POST":
        return https_fn.Response(
            json.dumps({"error": "Only POST requests are accepted"}),
            status=405,
            content_type="application/json"
        )

    try:
        data = req.get_json()
        if not data:
            return https_fn.Response(
                json.dumps({"error": "Missing JSON payload"}),
                status=400,
                content_type="application/json"
            )

        file_name = data.get("file_name")
        quality_flag = data.get("quality_flag")

        if not file_name or not isinstance(file_name, str):
            return https_fn.Response(
                json.dumps({"error": "Missing or invalid file_name"}),
                status=400,
                content_type="application/json"
            )

        if quality_flag is None or not isinstance(quality_flag, bool):
            return https_fn.Response(
                json.dumps({"error": "Missing or invalid quality_flag (must be boolean)"}),
                status=400,
                content_type="application/json"
            )

        from core.firebase_utils import _init_firebase
        _init_firebase()
        from firebase_admin import firestore
        db = firestore.client()

        # Determine the project ID from env, or use the fallback
        project_id = os.environ.get("FIREBASE_PROJECT_ID", "cartorio-meneghel-ai")

        doc_ref = db.collection("audit_logs").document()
        doc_ref.set({
            "file_name": file_name,
            "quality_flag": quality_flag,
            "timestamp": firestore.SERVER_TIMESTAMP,
            "project_id": project_id
        })

        return https_fn.Response(
            json.dumps({"status": "success", "message": "Audit event logged successfully"}),
            status=200,
            content_type="application/json"
        )
    except Exception as e:
        return https_fn.Response(
            json.dumps({"error": f"Internal server error: {str(e)}"}),
            status=500,
            content_type="application/json"
        )

@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.MB_256)
def log_hitl_resolution(req: https_fn.Request) -> https_fn.Response:
    """
    Logs an audit event for a Human-in-the-Loop discrepancy resolution.
    Accepts POST requests with JSON payload:
    {"document_id": "...", "field": "...", "expected": "...", "found": "...", "resolution_type": "..."}
    """
    if req.method != "POST":
        return https_fn.Response(
            json.dumps({"error": "Only POST requests are accepted"}),
            status=405,
            content_type="application/json"
        )

    try:
        data = req.get_json()
        if not data:
            return https_fn.Response(
                json.dumps({"error": "Missing JSON payload"}),
                status=400,
                content_type="application/json"
            )

        document_id = data.get("document_id")
        field = data.get("field")
        expected = data.get("expected")
        found = data.get("found")
        resolution_type = data.get("resolution_type")

        if not document_id or not isinstance(document_id, str):
            return https_fn.Response(
                json.dumps({"error": "Missing or invalid document_id"}),
                status=400,
                content_type="application/json"
            )

        if not resolution_type or not isinstance(resolution_type, str):
            return https_fn.Response(
                json.dumps({"error": "Missing or invalid resolution_type"}),
                status=400,
                content_type="application/json"
            )

        event_data = {
            "event_type": "hitl_resolution",
            "document_id": document_id,
            "field": field,
            "expected": expected,
            "found": found,
            "resolution_type": resolution_type,
            "project_id": os.environ.get("FIREBASE_PROJECT_ID", "cartorio-meneghel-ai")
        }

        # Asynchronously log the event to Firestore
        from core.audit import log_audit_event_async
        log_audit_event_async(event_data)

        return https_fn.Response(
            json.dumps({"status": "success", "message": "HitL resolution event logged successfully"}),
            status=200,
            content_type="application/json"
        )

    except Exception as e:
        logger.error("Error in log_hitl_resolution", exc_info=True)
        return https_fn.Response(
            json.dumps({
                "error": f"Internal server error: {str(e)}"
            }),
            status=500,
            content_type="application/json"
        )

@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.MB_512, timeout_sec=120)
def ingest_raw_clauses(req: https_fn.Request) -> https_fn.Response:
    """
    Ingests raw legal text, extracts clauses, vectorizes them, and saves to Firestore.
    Accepts POST requests with JSON payload: {"raw_text": "...", "cartorio_id": "..."}
    """
    if req.method != "POST":
        return https_fn.Response(json.dumps({"error": {"message": "Only POST requests are accepted", "status": "METHOD_NOT_ALLOWED"}}), status=405, content_type="application/json")

    auth_header = req.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        return https_fn.Response(json.dumps({"error": {"message": "Unauthenticated", "status": "UNAUTHENTICATED"}}), status=401, content_type="application/json")

    token = auth_header.split("Bearer ")[1]

    try:
        from core.firebase_utils import _init_firebase
        _init_firebase()
        from firebase_admin import auth, firestore
        from google.cloud.firestore_v1.vector import Vector

        try:
            decoded_token = auth.verify_id_token(token)
            uid = decoded_token.get("uid")
        except Exception as e:
            return https_fn.Response(json.dumps({"error": {"message": f"Invalid token: {str(e)}", "status": "UNAUTHENTICATED"}}), status=401, content_type="application/json")

        db = firestore.client()
        user_doc = db.collection('users').document(uid).get()
        if not user_doc.exists:
            return https_fn.Response(json.dumps({"error": {"message": "User not found", "status": "PERMISSION_DENIED"}}), status=403, content_type="application/json")

        user_data = user_doc.to_dict()
        user_role = user_data.get('role')
        user_cartorio_id = user_data.get('cartorio_id')

        if user_role not in ['super_admin']:
            return https_fn.Response(json.dumps({"error": {"message": "Unauthorized: Only super_admin can ingest clauses", "status": "PERMISSION_DENIED"}}), status=403, content_type="application/json")

        payload = req.get_json(silent=True)
        if not payload:
            return https_fn.Response(json.dumps({"error": {"message": "Missing JSON payload", "status": "INVALID_ARGUMENT"}}), status=400, content_type="application/json")

        raw_text = payload.get("raw_text")
        cartorio_id = payload.get("cartorio_id", "SYSTEM")

        if not raw_text:
            return https_fn.Response(json.dumps({"error": {"message": "Missing raw_text", "status": "INVALID_ARGUMENT"}}), status=400, content_type="application/json")

        from core.generator import parse_clause_with_llm, vectorize_text

        # 1. Parse Clauses
        parsed_data = parse_clause_with_llm(raw_text)
        clauses = parsed_data.get('clauses', [])

        if not clauses:
            return https_fn.Response(json.dumps({"error": {"message": "Failed to extract clauses from text", "status": "INTERNAL"}}), status=500, content_type="application/json")

        saved_clauses = []
        batch = db.batch()
        clauses_ref = db.collection("clauses")

        # 2. Vectorize and prepare batch
        for clause_def in clauses:
            title = clause_def.get("title", "")
            text = clause_def.get("text", "")

            # Combine title and text for rich embedding context
            embedding_input = f"Title: {title}\nText: {text}"
            embedding_vector = vectorize_text(embedding_input)

            if not embedding_vector:
                logger.warning(f"Skipping clause '{title}' due to embedding failure.")
                continue

            doc_ref = clauses_ref.document()

            clause_data = {
                "id": doc_ref.id,
                "title": title,
                "text": text,
                "required_variables": clause_def.get("required_variables", []),
                "scope_tags": clause_def.get("scope_tags", []),
                "embedding": Vector(embedding_vector),  # Important: use Firestore Vector type
                "is_active": True,
                "cartorio_id": cartorio_id,
                "created_at": firestore.SERVER_TIMESTAMP,
                "updated_at": firestore.SERVER_TIMESTAMP,
                "version": 1
            }
            batch.set(doc_ref, clause_data)

            # For response payload, we stringify the timestamp and omit the dense vector for brevity
            clause_data["embedding"] = "[VECTOR HIDDEN]"
            clause_data["created_at"] = "SERVER_TIMESTAMP"
            clause_data["updated_at"] = "SERVER_TIMESTAMP"
            saved_clauses.append(clause_data)

        # 3. Commit to Firestore
        batch.commit()

        return https_fn.Response(
            json.dumps({"data": {"status": "success", "ingested_clauses_count": len(saved_clauses), "clauses": saved_clauses}}),
            status=200,
            content_type="application/json"
        )

    except Exception as e:
        logger.error("Error in ingest_raw_clauses", exc_info=True)
        return https_fn.Response(json.dumps({"error": {"message": f"Internal server error: {str(e)}", "status": "INTERNAL"}}), status=500, content_type="application/json")

@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.MB_256)
def register_template(req: https_fn.Request) -> https_fn.Response:
    """
    Registers a new .docx template.
    Accepts Callable payload: {"data": {"cartorio_id": "...", "gcs_path": "...", "name": "...", "document_type": "...", "created_by": "..."}}
    Downloads the template from GCS, extracts Jinja2 tags, and creates a record in Firestore.
    """
    if req.method != "POST":
        return https_fn.Response(
            json.dumps({"error": {"message": "Only POST requests are accepted", "status": "METHOD_NOT_ALLOWED"}}),
            status=405,
            content_type="application/json"
        )

    auth_header = req.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        return https_fn.Response(
            json.dumps({"error": {"message": "Unauthenticated", "status": "UNAUTHENTICATED"}}),
            status=401,
            content_type="application/json"
        )

    token = auth_header.split("Bearer ")[1]

    try:
        from core.firebase_utils import _init_firebase
        _init_firebase()
        from firebase_admin import auth, firestore, storage

        try:
            decoded_token = auth.verify_id_token(token)
            uid = decoded_token.get("uid")
        except Exception as e:
            return https_fn.Response(
                json.dumps({"error": {"message": f"Invalid token: {str(e)}", "status": "UNAUTHENTICATED"}}),
                status=401,
                content_type="application/json"
            )

        db = firestore.client()

        user_doc = db.collection('users').document(uid).get()
        if not user_doc.exists:
            return https_fn.Response(
                json.dumps({"error": {"message": "User not found", "status": "PERMISSION_DENIED"}}),
                status=403,
                content_type="application/json"
            )

        user_data = user_doc.to_dict()

        payload = req.get_json(silent=True)
        if not payload:
            return https_fn.Response(
                json.dumps({"error": {"message": "Missing JSON payload", "status": "INVALID_ARGUMENT"}}),
                status=400,
                content_type="application/json"
            )

        data = payload.get("data", {})

        cartorio_id = data.get("cartorio_id")
        user_role = user_data.get('role')
        user_cartorio_id = user_data.get('cartorio_id')

        if not cartorio_id:
            return https_fn.Response(
                json.dumps({"error": {"message": "Unauthorized: Missing cartorio_id", "status": "PERMISSION_DENIED"}}),
                status=403,
                content_type="application/json"
            )

        if user_role not in ['cartorio_admin', 'super_admin']:
            return https_fn.Response(
                json.dumps({"error": {"message": "Unauthorized: Role not permitted", "status": "PERMISSION_DENIED"}}),
                status=403,
                content_type="application/json"
            )

        if user_role != 'super_admin' and user_cartorio_id != cartorio_id:
            return https_fn.Response(
                json.dumps({"error": {"message": "Unauthorized: Cartorio mismatch", "status": "PERMISSION_DENIED"}}),
                status=403,
                content_type="application/json"
            )

        cartorio_id = data.get("cartorio_id")
        gcs_path = data.get("gcs_path")
        name = data.get("name")
        document_type = data.get("document_type")
        created_by = data.get("created_by")

        if not all([cartorio_id, gcs_path, name, document_type, created_by]):
            return https_fn.Response(
                json.dumps({"error": {"message": "Missing required fields", "status": "INVALID_ARGUMENT"}}),
                status=400,
                content_type="application/json"
            )

        bucket = storage.bucket()

        blob = bucket.blob(gcs_path)
        if not blob.exists():
            return https_fn.Response(
                json.dumps({"error": {"message": f"File not found at {gcs_path}", "status": "NOT_FOUND"}}),
                status=404,
                content_type="application/json"
            )

        template_bytes = blob.download_as_bytes()

        from core.generator import extract_tags_from_template, generate_roles_schema_for_template
        try:
            required_tags = extract_tags_from_template(template_bytes)
        except ValueError as ve:
            return https_fn.Response(
                json.dumps({"error": {"message": str(ve), "status": "INVALID_ARGUMENT"}}),
                status=400,
                content_type="application/json"
            )

        try:
            roles_schema = generate_roles_schema_for_template(required_tags)
        except Exception as e:
            logger.error(f"Error generating roles schema for template: {e}")
            roles_schema = []

        doc_ref = db.collection("templates").document()

        template_data = {
            "id": doc_ref.id,
            "cartorio_id": cartorio_id,
            "name": name,
            "document_type": document_type,
            "gcs_path": gcs_path,
            "required_tags": required_tags,
            "roles_schema": roles_schema,
            "created_by": created_by,
            "created_at": firestore.SERVER_TIMESTAMP,
            "is_active": True
        }

        doc_ref.set(template_data)

        template_data["created_at"] = "SERVER_TIMESTAMP"

        return https_fn.Response(
            json.dumps({"data": {"status": "success", "template": template_data}}),
            status=200,
            content_type="application/json"
        )

    except Exception as e:
        logger.error("Error in register_template", exc_info=True)
        return https_fn.Response(
            json.dumps({"error": {"message": f"Internal server error: {str(e)}", "status": "INTERNAL"}}),
            status=500,
            content_type="application/json"
        )

@https_fn.on_call(memory=options.MemoryOption.MB_256)
def grantSupportAccess(req: https_fn.CallableRequest) -> dict:
    """
    Grants time-bound Break-Glass support access for a specific document to super_admins.
    Accepts Callable payload: {"document_id": "...", "duration_hours": int}
    """
    if not req.auth:
        raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.UNAUTHENTICATED, message="Unauthenticated")

    try:
        from core.firebase_utils import _init_firebase
        _init_firebase()
        from firebase_admin import firestore
        import datetime

        db = firestore.client()

        user_doc = db.collection('users').document(req.auth.uid).get()
        if not user_doc.exists:
            raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.PERMISSION_DENIED, message="User not found")

        user_data = user_doc.to_dict()
        data = req.data

        document_id = data.get("document_id")
        duration_hours = data.get("duration_hours", 24)

        if not document_id:
            raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.INVALID_ARGUMENT, message="Missing document_id")

        if user_data.get('role') != 'cartorio_admin':
            raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.PERMISSION_DENIED, message="Only cartorio_admin can grant support access")

        minuta_ref = db.collection('minutas').document(document_id)
        minuta_doc = minuta_ref.get()

        if not minuta_doc.exists:
            raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.NOT_FOUND, message="Document not found")

        minuta_data = minuta_doc.to_dict()
        if minuta_data.get('cartorio_id') != user_data.get('cartorio_id'):
            raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.PERMISSION_DENIED, message="Document does not belong to your cartorio")

        expires_at = datetime.datetime.now(datetime.timezone.utc) + datetime.timedelta(hours=duration_hours)

        minuta_ref.update({
            "support_grant": {
                "expires_at": expires_at,
                "granted_by": req.auth.uid
            }
        })

        from core.audit import log_audit_event_async
        event_data = {
            "event_type": "support_access_granted",
            "document_id": document_id,
            "cartorio_id": user_data.get('cartorio_id'),
            "duration_hours": duration_hours,
            "granted_by": req.auth.uid,
            "project_id": os.environ.get("FIREBASE_PROJECT_ID", "cartorio-meneghel-ai")
        }
        log_audit_event_async(event_data)

        return {"status": "success", "message": f"Support access granted for {duration_hours} hours"}

    except https_fn.HttpsError as he:
        raise he
    except Exception as e:
        logger.error("Error in grantSupportAccess", exc_info=True)
        raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.INTERNAL, message=f"Internal server error: {str(e)}")

@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.MB_512, timeout_sec=540)
def generate_document_api(req: https_fn.Request) -> https_fn.Response:
    """
    Generates a document from a registered template.
    Accepts Callable payload: {"cartorio_id": "...", "template_id": "...", "verified_data": {...}}
    """
    if req.method != "POST":
        return https_fn.Response(json.dumps({"error": {"code": "METHOD_NOT_ALLOWED", "message": "Method not allowed"}}), status=405, content_type="application/json")

    auth_header = req.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        return https_fn.Response(json.dumps({"error": {"code": "UNAUTHENTICATED", "message": "Unauthenticated"}}), status=401, content_type="application/json")

    token = auth_header.split("Bearer ")[1]

    try:
        from core.firebase_utils import _init_firebase
        _init_firebase()
        from firebase_admin import auth, firestore, storage
        db = firestore.client()

        try:
            decoded_token = auth.verify_id_token(token)
            uid = decoded_token.get("uid")
        except Exception:
            return https_fn.Response(json.dumps({"error": {"code": "UNAUTHENTICATED", "message": "Invalid token"}}), status=401, content_type="application/json")

        user_doc = db.collection('users').document(uid).get()
        if not user_doc.exists:
            raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.PERMISSION_DENIED, message="User not found")

        user_data = user_doc.to_dict()
        data = req.get_json(silent=True)

        if not data:
            return https_fn.Response(
                json.dumps({"error": {"code": "INVALID_ARGUMENT", "message": "Missing or malformed JSON payload"}}),
                status=400,
                content_type="application/json"
            )

        cartorio_id = data.get("cartorio_id")
        user_role = user_data.get('role')
        user_cartorio_id = user_data.get('cartorio_id')

        if not cartorio_id:
            raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.PERMISSION_DENIED, message="Unauthorized")

        if user_role != 'super_admin' and user_cartorio_id != cartorio_id:
            raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.PERMISSION_DENIED, message="Unauthorized")


        cartorio_id = data.get("cartorio_id")
        template_id = data.get("template_id")
        verified_data = data.get("verified_data")
        draft_id = data.get("draft_id")
        imported_at = data.get("imported_at")
        role_mapping = data.get("role_mapping", {})

        if not cartorio_id or not template_id or verified_data is None:
            raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.INVALID_ARGUMENT, message="Missing required fields")

        if draft_id and imported_at:
            # Check for optimistic concurrency
            minuta_ref = db.collection("minutas").document(draft_id)
            minuta_doc = minuta_ref.get()
            if minuta_doc.exists:
                minuta_data = minuta_doc.to_dict()
                if minuta_data.get('cartorio_id') != cartorio_id:
                    raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.PERMISSION_DENIED, message="Unauthorized")

                db_updated_at = minuta_data.get("updatedAt")
                if db_updated_at:
                    # Compare timestamps
                    # FireStore returns datetime objects for timestamps
                    from datetime import datetime, timezone

                    try:
                        imported_seconds = imported_at.get('_seconds', 0)
                        imported_nanos = imported_at.get('_nanoseconds', 0)
                        imported_dt = datetime.fromtimestamp(imported_seconds + imported_nanos / 1e9, tz=timezone.utc)

                        # Firestore timestamp objects have a timestamp() method, or can be converted.
                        # Just to be safe, extract seconds/nanos from db_updated_at or use timestamp()
                        if hasattr(db_updated_at, 'timestamp'):
                            db_ts = db_updated_at.timestamp()
                        else:
                            # Assume it's a datetime
                            db_ts = db_updated_at.timestamp()

                        imported_ts = imported_dt.timestamp()

                        is_stale = db_ts > imported_ts + 0.001
                    except Exception as ts_err:
                        # Log but don't crash if timestamp parsing fails for some reason
                        logger.error(f"Error parsing timestamps for concurrency check: {ts_err}")
                        is_stale = False

                    if is_stale:
                        raise https_fn.HttpsError(
                            code=https_fn.FunctionsErrorCode.FAILED_PRECONDITION,
                            message="Race condition detected: Minuta was modified after it was imported. Please re-import the latest data."
                        )


        import uuid
        selected_clause_ids = data.get("selected_clause_ids", [])

        if template_id == "DYNAMIC_CLAUSES":
            # Assembly template dynamically
            from core.generator import assemble_dynamic_document
            try:
                # We fetch ground_truth if we have draft_id
                ground_truth = {}
                if draft_id:
                     draft_doc = db.collection("minutas").document(draft_id).get()
                     if draft_doc.exists:
                         ground_truth = draft_doc.to_dict()
                generated_bytes = assemble_dynamic_document(selected_clause_ids, role_mapping, ground_truth, verified_data, db)
            except ValueError as ve:
                raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.INVALID_ARGUMENT, message=str(ve))
        else:
            bucket = storage.bucket()

            template_ref = db.collection("templates").document(template_id)
            template_doc = template_ref.get()

            if not template_doc.exists:
                raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.NOT_FOUND, message="Template not found")

            template_info = template_doc.to_dict()

            template_cartorio_id = template_info.get("cartorio_id")
            if template_cartorio_id not in [cartorio_id, "SYSTEM"]:
                raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.PERMISSION_DENIED, message="Unauthorized to access this template")

            gcs_path = template_info.get("gcs_path")
            required_tags = template_info.get("required_tags", [])

            blob = bucket.blob(gcs_path)
            if not blob.exists():
                raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.NOT_FOUND, message="Template file missing in storage")

            template_bytes = blob.download_as_bytes()

            from core.generator import generate_document_from_template
            try:
                generated_bytes = generate_document_from_template(template_bytes, verified_data, required_tags)
            except ValueError as ve:
                raise https_fn.HttpsError(code=https_fn.FunctionsErrorCode.INVALID_ARGUMENT, message=str(ve))

        import base64
        base64_encoded = base64.b64encode(generated_bytes).decode('utf-8')

        from docx import Document
        import io
        try:
            doc_parsed = Document(io.BytesIO(generated_bytes))
            plain_text = '\n'.join([p.text for p in doc_parsed.paragraphs])
        except Exception as e:
            logger.error(f"Failed to parse docx for plain text: {e}")
            plain_text = ""

        return https_fn.Response(
            json.dumps({"status": "success", "file_base64": base64_encoded, "plain_text": plain_text}),
            status=200,
            content_type="application/json"
        )

    except https_fn.HttpsError as he:
        # Expected error contract for frontend
        error_payload = {
            "error": {
                "code": he.code.name,
                "message": he.message
            }
        }
        # HttpsError typically returns 500 if we raise it in on_request natively unless mapped,
        # but since we are doing on_request, let's map standard error codes to HTTP statuses
        status_map = {
            https_fn.FunctionsErrorCode.INVALID_ARGUMENT: 400,
            https_fn.FunctionsErrorCode.UNAUTHENTICATED: 401,
            https_fn.FunctionsErrorCode.PERMISSION_DENIED: 403,
            https_fn.FunctionsErrorCode.NOT_FOUND: 404,
            https_fn.FunctionsErrorCode.FAILED_PRECONDITION: 400,
            https_fn.FunctionsErrorCode.INTERNAL: 500,
        }
        http_status = status_map.get(he.code, 500)
        return https_fn.Response(json.dumps(error_payload), status=http_status, content_type="application/json")
    except Exception as e:
        logger.error("Error in generate_document_api", exc_info=True)
        return https_fn.Response(
            json.dumps({"error": {"code": "INTERNAL", "message": f"Internal server error: {str(e)}"}}),
            status=500,
            content_type="application/json"
        )

@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.MB_512, timeout_sec=540)
def preview_dynamic_document(req: https_fn.Request) -> https_fn.Response:
    """
    Generates a plain text preview of the dynamic document, ensuring the master wrapper is applied.
    Accepts REST payload: {"cartorio_id": "...", "verified_data": {...}, "role_mapping": {...}, "selected_clause_ids": [...], "draft_id": "..."}
    """
    if req.method != "POST":
        return https_fn.Response(json.dumps({"error": {"code": "METHOD_NOT_ALLOWED", "message": "Method not allowed"}}), status=405, content_type="application/json")

    auth_header = req.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        return https_fn.Response(json.dumps({"error": {"code": "UNAUTHENTICATED", "message": "Unauthenticated"}}), status=401, content_type="application/json")

    token = auth_header.split("Bearer ")[1]

    try:
        from core.firebase_utils import _init_firebase
        _init_firebase()
        from firebase_admin import auth, firestore
        db = firestore.client()

        try:
            decoded_token = auth.verify_id_token(token)
            uid = decoded_token.get("uid")
        except Exception:
            return https_fn.Response(json.dumps({"error": {"code": "UNAUTHENTICATED", "message": "Invalid token"}}), status=401, content_type="application/json")

        user_doc = db.collection('users').document(uid).get()
        if not user_doc.exists:
            return https_fn.Response(json.dumps({"error": {"code": "PERMISSION_DENIED", "message": "User not found"}}), status=403, content_type="application/json")

        user_data = user_doc.to_dict()
        data = req.get_json(silent=True)

        if not data:
            return https_fn.Response(json.dumps({"error": {"code": "INVALID_ARGUMENT", "message": "Missing JSON payload"}}), status=400, content_type="application/json")

        cartorio_id = data.get("cartorio_id")
        user_role = user_data.get('role')
        user_cartorio_id = user_data.get('cartorio_id')

        if not cartorio_id:
            return https_fn.Response(json.dumps({"error": {"code": "PERMISSION_DENIED", "message": "Unauthorized"}}), status=403, content_type="application/json")

        if user_role != 'super_admin' and user_cartorio_id != cartorio_id:
            return https_fn.Response(json.dumps({"error": {"code": "PERMISSION_DENIED", "message": "Unauthorized"}}), status=403, content_type="application/json")

        verified_data = data.get("verified_data", {})
        draft_id = data.get("draft_id")
        role_mapping = data.get("role_mapping", {})
        selected_clause_ids = data.get("selected_clause_ids", [])

        from core.generator import assemble_dynamic_document
        ground_truth = {}
        if draft_id:
             draft_doc = db.collection("minutas").document(draft_id).get()
             if draft_doc.exists:
                 ground_truth = draft_doc.to_dict()

        try:
            generated_bytes = assemble_dynamic_document(selected_clause_ids, role_mapping, ground_truth, verified_data, db)
        except ValueError as ve:
            return https_fn.Response(json.dumps({"error": {"code": "INVALID_ARGUMENT", "message": str(ve)}}), status=400, content_type="application/json")

        from docx import Document
        import io
        try:
            doc_parsed = Document(io.BytesIO(generated_bytes))
            plain_text = '\n'.join([p.text for p in doc_parsed.paragraphs])
        except Exception as e:
            logger.error(f"Failed to parse docx for plain text preview: {e}")
            plain_text = ""

        return https_fn.Response(
            json.dumps({"status": "success", "text": plain_text}),
            status=200,
            content_type="application/json"
        )

    except Exception as e:
        logger.error("Error in preview_dynamic_document", exc_info=True)
        return https_fn.Response(
            json.dumps({"error": {"code": "INTERNAL", "message": f"Internal server error: {str(e)}"}}),
            status=500,
            content_type="application/json"
        )

@https_fn.on_request(cors=global_cors, memory=options.MemoryOption.MB_256)
def suggest_field_text(req: https_fn.Request) -> https_fn.Response:
    """
    Suggests text for a specific form field using LLM and context data.
    Accepts REST payload: {"cartorio_id": "...", "tag": "...", "context_data": {...}}
    """
    if req.method != "POST":
        return https_fn.Response(json.dumps({"error": {"code": "METHOD_NOT_ALLOWED", "message": "Method not allowed"}}), status=405, content_type="application/json")

    auth_header = req.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        return https_fn.Response(json.dumps({"error": {"code": "UNAUTHENTICATED", "message": "Unauthenticated"}}), status=401, content_type="application/json")

    token = auth_header.split("Bearer ")[1]

    try:
        from core.firebase_utils import _init_firebase
        _init_firebase()
        from firebase_admin import auth, firestore

        try:
            decoded_token = auth.verify_id_token(token)
            user_role = decoded_token.get("role")
            user_cartorio = decoded_token.get("cartorio_id")
        except Exception as e:
            return https_fn.Response(json.dumps({"error": {"code": "UNAUTHENTICATED", "message": "Invalid token"}}), status=401, content_type="application/json")

        try:
            req_data = req.get_json(silent=True) or {}
            cartorio_id = req_data.get("cartorio_id")
            tag = req_data.get("tag")
            context_data = req_data.get("context_data")

            if not cartorio_id or not tag or context_data is None:
                return https_fn.Response(json.dumps({"error": {"code": "INVALID_ARGUMENT", "message": "Missing required fields"}}), status=400, content_type="application/json")

            if user_role != "super_admin" and cartorio_id != user_cartorio:
                return https_fn.Response(json.dumps({"error": {"code": "PERMISSION_DENIED", "message": "Tenant mismatch"}}), status=403, content_type="application/json")

            from core.generator import suggest_field_text_llm
            suggestion = suggest_field_text_llm(tag, context_data)

            return https_fn.Response(json.dumps({"status": "success", "suggestion": suggestion}), status=200, content_type="application/json")

        except Exception as e:
            import traceback
            traceback.print_exc()
            return https_fn.Response(json.dumps({"error": {"code": "INTERNAL", "message": f"Server error: {str(e)}"}}), status=500, content_type="application/json")

    except Exception as e:
        return https_fn.Response(json.dumps({"error": {"code": "INTERNAL", "message": "Initialization error"}}), status=500, content_type="application/json")
