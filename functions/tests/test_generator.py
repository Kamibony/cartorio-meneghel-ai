import pytest
from core.generator import extract_tags_from_template

def test_extract_tags():
    from docx import Document
    import io

    d = Document()
    d.add_paragraph("Hello {{ name }}, you are {{ age }} years old.")
    f = io.BytesIO()
    d.save(f)

    tags = extract_tags_from_template(f.getvalue())
    assert set(tags) == {"name", "age"}

def test_is_literal_tag():
    from core.generator import is_literal_tag
    assert is_literal_tag("valor_imovel") == True
    assert is_literal_tag("VALOR_TOTAL") == True
    assert is_literal_tag("emolumentos") == True
    assert is_literal_tag("taxa_emolumentos") == True
    assert is_literal_tag("nome_comprador") == False
    assert is_literal_tag("data_nascimento") == False

def test_generate_document_from_template_mocked_llm(monkeypatch):
    import io
    from docx import Document
    from core.generator import generate_document_from_template

    d = Document()
    d.add_paragraph("Comprador: {{ nome_comprador }}")
    d.add_paragraph("Valor: {{ valor_imovel }}")
    f = io.BytesIO()
    d.save(f)
    template_bytes = f.getvalue()

    verified_data = {
        "nome_comprador": "João da Silva",
        "valor_imovel": "R$ 500.000,00"
    }
    required_tags = ["nome_comprador", "valor_imovel"]

    class MockModelResponse:
        @property
        def text(self):
            return '{"nome_comprador": "Sr. João da Silva"}'

    class MockModels:
        def generate_content(self, **kwargs):
            return MockModelResponse()

    class MockClient:
        def __init__(self, **kwargs):
            self.models = MockModels()

    import google.genai as genai
    monkeypatch.setattr(genai, "Client", MockClient)

    # Allow avoiding real network calls
    monkeypatch.setenv("FIREBASE_PROJECT_ID", "test")
    monkeypatch.setenv("VERTEX_AI_LOCATION", "test")

    generated_bytes = generate_document_from_template(template_bytes, verified_data, required_tags)

    assert generated_bytes is not None

    # Load the generated document and verify contents
    from docxtpl import DocxTemplate
    doc = DocxTemplate(io.BytesIO(generated_bytes))

    # We can't easily parse the resulting docx cleanly without python-docx since docxtpl evaluates it,
    # but we can check if it rendered properly without crashing.
    import zipfile
    zf = zipfile.ZipFile(io.BytesIO(generated_bytes))
    xml_content = zf.read("word/document.xml").decode("utf-8")

    # Check if literal tags and grammar tags were applied
    assert "Sr. João da Silva" in xml_content # from mocked LLM
    assert "R$ 500.000,00" in xml_content # from literal tag fallback
