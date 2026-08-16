import re
import logging

logger = logging.getLogger(__name__)

class DocumentResolver:
    def __init__(self):
        pass

    def get_entity_by_id(self, entity_id: str, ground_truth: dict) -> dict:
        entities = ground_truth.get("entities", [])
        if not entities and "_contexto_extraido" in ground_truth:
            entities = ground_truth.get("_contexto_extraido", {}).get("entities", [])
        for ent in entities:
            if ent.get("id") == entity_id:
                return ent
            # Fallback to matching by name if explicit UUID fails (e.g. LLM generated names or frontend fallback ID)
            ent_name = ent.get("nome") or ent.get("razao_social") or ent.get("name") or ent.get("entity_name")
            if ent_name and ent_name == entity_id:
                return ent
        return {}

    def get_entity_value(self, entity: dict, attr_name: str) -> str:
        attr_name = attr_name.lower()
        if attr_name in ["nome", "razao_social", "name"]:
            return entity.get("nome") or entity.get("razao_social") or entity.get("name") or entity.get("entity_name") or ""

        val = entity.get(attr_name)
        if val:
            return val

        for attr in entity.get("attributes", []):
            if attr.get("key", "").lower() == attr_name:
                return attr.get("value", "")
        return ""

    def generate_preamble(self, role_mapping: dict, ground_truth: dict) -> str:
        """
        Dynamically generates the 'Qualificação' section.
        """
        preamble_parts = []

        # Hardcode specific logic for expected roles
        for role, entity_ids in role_mapping.items():
            if not entity_ids:
                continue

            entities = [self.get_entity_by_id(eid, ground_truth) for eid in entity_ids]

            # Filter out empty
            entities = [e for e in entities if e]
            if not entities:
                continue

            role_title = role.upper()
            if len(entities) > 1:
                role_title += "S" # Simple pluralization

            entity_descriptions = []
            for e in entities:
                nome = self.get_entity_value(e, "nome")
                nacionalidade = self.get_entity_value(e, "nacionalidade")
                estado_civil = self.get_entity_value(e, "estado_civil")
                profissao = self.get_entity_value(e, "profissao")
                cpf = self.get_entity_value(e, "cpf") or self.get_entity_value(e, "cnpj")
                rg = self.get_entity_value(e, "rg")
                endereco = self.get_entity_value(e, "endereco")

                desc = f"{nome}"
                if nacionalidade: desc += f", {nacionalidade}"
                if estado_civil: desc += f", {estado_civil}"
                if profissao: desc += f", {profissao}"
                if rg: desc += f", portador do RG {rg}"
                if cpf: desc += f", inscrito no CPF/CNPJ sob o nº {cpf}"
                if endereco: desc += f", residente e domiciliado em {endereco}"

                entity_descriptions.append(desc)

            if len(entity_descriptions) > 1:
                combined_desc = ", ".join(entity_descriptions[:-1]) + " e " + entity_descriptions[-1]
            else:
                combined_desc = entity_descriptions[0]

            preamble_parts.append(f"{role_title}: {combined_desc}.")

        if not preamble_parts:
            return "QUALIFICAÇÃO: [DADOS FALTANTES]"

        return "QUALIFICAÇÃO:\n\n" + "\n".join(preamble_parts)

    def resolve_tag(self, tag: str, role_mapping: dict, ground_truth: dict, variables_data: dict) -> str:
        """
        Parses tags and retrieves exact data deterministically.
        """
        import json
        # Try finding in explicit variables_data first
        if tag in variables_data and variables_data[tag]:
            val = variables_data[tag]
            if isinstance(val, (dict, list)):
                return json.dumps(val, ensure_ascii=False)
            return str(val)

        # Try parsing from role mapping e.g., OUTORGANTE_NOME
        parts = tag.split("_", 1)
        if len(parts) == 2:
            role = parts[0].upper()
            attr = parts[1].lower()

            if role in role_mapping:
                entity_ids = role_mapping[role]
                if entity_ids:
                    entities = [self.get_entity_by_id(eid, ground_truth) for eid in entity_ids]
                    entities = [e for e in entities if e]

                    values = [self.get_entity_value(e, attr) for e in entities]
                    values = [v for v in values if v]

                    if values:
                        if len(values) > 1:
                            return ", ".join(values[:-1]) + " e " + values[-1]
                        return values[0]

        return "[DADO FALTANTE]"

    def _get_role_qualification(self, role: str, role_mapping: dict, ground_truth: dict) -> str:
        entity_ids = role_mapping.get(role) or role_mapping.get(role.upper()) or role_mapping.get(role.lower())
        if not entity_ids:
            return "[DADO FALTANTE]"

        entities = [self.get_entity_by_id(eid, ground_truth) for eid in entity_ids]
        entities = [e for e in entities if e]
        if not entities:
            return "[DADO FALTANTE]"

        entity_descriptions = []
        for e in entities:
            nome = self.get_entity_value(e, "nome")
            nacionalidade = self.get_entity_value(e, "nacionalidade")
            estado_civil = self.get_entity_value(e, "estado_civil")
            profissao = self.get_entity_value(e, "profissao")
            cpf = self.get_entity_value(e, "cpf") or self.get_entity_value(e, "cnpj")
            rg = self.get_entity_value(e, "rg")
            endereco = self.get_entity_value(e, "endereco")

            desc = f"{nome}"
            if nacionalidade: desc += f", {nacionalidade}"
            if estado_civil: desc += f", {estado_civil}"
            if profissao: desc += f", {profissao}"
            if rg: desc += f", portador do RG {rg}"
            if cpf: desc += f", inscrito no CPF/CNPJ sob o nº {cpf}"
            if endereco: desc += f", residente e domiciliado em {endereco}"

            entity_descriptions.append(desc)

        if len(entity_descriptions) > 1:
            combined_desc = ", ".join(entity_descriptions[:-1]) + " e " + entity_descriptions[-1]
        else:
            combined_desc = entity_descriptions[0]

        return combined_desc

    def assemble(self, selected_clauses: list, role_mapping: dict, ground_truth: dict, variables_data: dict) -> bytes:
        from docx import Document
        import io

        qualificacao_outorgante = self._get_role_qualification("OUTORGANTE", role_mapping, ground_truth)
        qualificacao_procurador = self._get_role_qualification("PROCURADOR", role_mapping, ground_truth)

        poderes_parts = []
        for clause_data in selected_clauses:
            title = clause_data.get("title", "")
            text = clause_data.get("text", "")

            # Find all tags in the text
            found_tags = set(re.findall(r"\{\{([A-Za-z0-9_]+)\}\}", text))

            # Resolve tags
            for tag in found_tags:
                resolved_value = self.resolve_tag(tag, role_mapping, ground_truth, variables_data)
                text = text.replace(f"{{{{{tag}}}}}", str(resolved_value))

            if title:
                poderes_parts.append(f"{title.upper()}\n{text}")
            else:
                poderes_parts.append(text)

        poderes_especificos = "\n\n".join(poderes_parts)

        master_template = f'''MINUTA DE PROCURAÇÃO PÚBLICA

OUTORGANTE(S):
{qualificacao_outorgante}

OUTORGADO(A)(S) / PROCURADOR(A)(ES):
{qualificacao_procurador}

PODERES:
Por este instrumento público e nos melhores termos de direito, o(s) Outorgante(s) nomeia(m) e constitui(em) seu(s) bastante Procurador(es) acima qualificado(s), a quem confere(m) os mais amplos, gerais e ilimitados poderes para o fim específico de:

{poderes_especificos}

Podendo, para tanto, assinar recibos, dar quitações, requerer, alegar e assinar o que for preciso, combinar cláusulas e condições, juntar e retirar documentos, prestar declarações e, enfim, praticar todos os demais atos necessários ao fiel e cabal cumprimento do presente mandato, comprometendo-se o(s) Outorgante(s) a dar tudo por bom, firme e valioso.'''

        document = Document()
        for paragraph_text in master_template.split('\n'):
            if paragraph_text == "MINUTA DE PROCURAÇÃO PÚBLICA":
                document.add_heading(paragraph_text, level=0)
            elif paragraph_text.strip():
                document.add_paragraph(paragraph_text)
            else:
                document.add_paragraph()

        out_f = io.BytesIO()
        document.save(out_f)
        return out_f.getvalue()
